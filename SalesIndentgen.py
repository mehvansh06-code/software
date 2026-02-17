import customtkinter as ctk
import os
import sys
import logging
import subprocess
import threading
import json
import time
from datetime import date, datetime
from tkinter import messagebox
from tkcalendar import DateEntry
import openpyxl
from openpyxl import Workbook

# =================================================
# CRITICAL FIX FOR "NoneType has no attribute write"
# =================================================
if sys.stdout is None:
    sys.stdout = open(os.devnull, "w")
if sys.stderr is None:
    sys.stderr = open(os.devnull, "w")

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from num2words import num2words

# Try importing docx2pdf for PDF generation
try:
    from docx2pdf import convert as convert_to_pdf
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# =================================================
# CONFIGURATION & CONSTANTS
# =================================================
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

ctk.set_appearance_mode("Light")
ctk.set_default_color_theme("dark-blue")

# UI styling constants (visual only)
CARD_RADIUS = 10
CARD_BORDER = "#E0E0E0"
CARD_PAD = 12
SECTION_HEADER_FONT = ("Arial", 13, "bold")
LABEL_GRAY = "gray50"

# Determine App Directory (Frozen vs Script)
if getattr(sys, "frozen", False):
    APP_DIR = os.path.dirname(sys.executable)
else:
    APP_DIR = os.path.dirname(os.path.abspath(__file__))

DATA_DIR = os.path.join(APP_DIR, "data")
OUTPUT_DIR = os.path.join(APP_DIR, "Generated_Indents")

os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ----------------- BUSINESS RULES -----------------
GST_RATE_IGST = 0.05   # 5%
GST_RATE_CGST = 0.025  # 2.5%
GST_RATE_SGST = 0.025  # 2.5%

# ----------------- UI PLACEHOLDERS ----------------
PH_COMPANY = "Select Company..."
PH_BUYER = "Select Buyer..."
PH_PRODUCT = "Select Quality..."
PH_SHIP = "Select Consignee..."

PAYMENT_TERMS = [
    "Payment Fully Advance",
    "Payment Before Delivery",
    "Payment Against Delivery",
    "Payment Within 7 Days",
    "Payment Within 15 Days",
    "Payment Within 30 Days",
    "Payment Within 45 Days",
    "Payment Within 60 Days",
]

INCO_TERMS = ["CIF", "FOB", "EXW", "CFR", "DAP"]
EXPORT_CURRENCIES = ["USD", "GBP"]

# ----------------- COMPANY DATABASE ----------------
# Added 'iec' field here for Export
COMPANY_DB = {
    "Gujarat Flotex Pvt. Ltd.": {
        "address": "3rd Floor, Elanza Vertex, Behind Armieda, Sindhu Bhavan Road,\n"
                   "Pakwan Cross Road, Ahmedabad-380059, Gujarat (India)",
        "gstin": "24AABCG4542P1ZF",
        "iec": "0801010128",  # <--- UPDATE THIS WITH REAL IEC
        "phone": "6358858231",
        "bank_details": {
            "Account Holder": "GUJARAT FLOTEX PVT LTD",
            "Bank": "STATE BANK OF INDIA",
            "Branch": "LAGHU UDHYOG, AHMEDABAD",
            "Acct": "30852691460",
            "IFSC": "SBIN0003993",
            "SWIFT": "SBININBBA23",
        },
    },
    "GTEX Fabrics": {
        "address": "3rd Floor, Elanza Vertex, Sindhu Bhavan Road,\n"
                   "Ahmedabad - 380054, Gujarat (India)",
        "gstin": "24AAGCG4275J1ZG",
        "iec": "AAGCG4275J", # <--- UPDATE THIS WITH REAL IEC
        "phone": "6358858231",
        "bank_details": {
           "Account Holder": "GTEX FABRICS PVT LTD",
           "Bank": "STATE BANK OF INDIA",
            "Branch": "LAGHU UDHYOG, AHMEDABAD",
            "Acct": "39092267695",
            "IFSC": "SBIN0003993",
            "SWIFT": "SBININBBA23",
        },
    },
}


# =================================================
# HELPER FUNCTIONS
# =================================================
def open_file(path):
    """Platform-independent file opener."""
    try:
        if os.name == 'nt':  # Windows
            os.startfile(path)
        elif sys.platform == 'darwin':  # macOS
            subprocess.call(['open', path])
        else:  # Linux
            subprocess.call(['xdg-open', path])
    except Exception as e:
        logging.error(f"Could not open file: {e}")
        messagebox.showerror("Error", f"Could not open file automatically.\nPath: {path}")


# =================================================
# DATA MANAGER
# =================================================
class DataManager:
    products = []
    products_map = {} 
    
    # Split Customers
    customers_domestic = {}
    customers_export = {}

    @staticmethod
    def normalize(x):
        return str(x).strip().lower() if x is not None else ""

    @staticmethod
    def _find_col(headers: dict, candidates):
        for c in candidates:
            c_norm = DataManager.normalize(c)
            if c_norm in headers:
                return headers[c_norm]
        return None

    @staticmethod
    def _auto_split_master_file(master_path):
        """Automatically splits customers.xlsx into domestic and export files."""
        try:
            logging.info("Auto-splitting master customer file...")
            wb = openpyxl.load_workbook(master_path, data_only=True)
            ws = wb.active

            # Identify Headers
            headers = {DataManager.normalize(cell.value): i for i, cell in enumerate(ws[1]) if cell.value}
            
            # Key Columns
            col_country = DataManager._find_col(headers, ["Country", "Destination Country"])
            if col_country is None:
                logging.error("Could not find 'Country' column in master file. Cannot split.")
                return

            # Prepare New Workbooks
            wb_dom = Workbook()
            ws_dom = wb_dom.active
            wb_exp = Workbook()
            ws_exp = wb_exp.active

            # Headers for New Files
            headers_dom = ["Customer Name", "Billing Address", "State", "GST No", "Mobile", "Ship Site Name", "Shipping Address", "Sales Person Name", "Sales Person Mobile", "Sales Person Email", "Payment Terms"]
            headers_exp = ["Customer Name", "Billing Address", "Destination Country", "Mobile", "Consignee Name", "Consignee Address", "Sales Person Name", "Sales Person Mobile", "Sales Person Email", "Port of Loading", "Port of Discharge", "Payment Terms"]

            ws_dom.append(headers_dom)
            ws_exp.append(headers_exp)

            # Process Rows
            for row in ws.iter_rows(min_row=2, values_only=True):
                # Helper to safely get value
                def get_val(col_names):
                    idx = DataManager._find_col(headers, col_names)
                    if idx is not None and idx < len(row):
                        val = row[idx]
                        return str(val).strip() if val is not None else ""
                    return ""

                country = get_val(["Country", "Destination Country"])
                if not country: continue # Skip empty rows

                # Common Data
                c_name = get_val(["Customer Name", "Name"])
                c_bill = get_val(["Billing Address"])
                c_mob = get_val(["Mobile", "Phone"])
                c_sp = get_val(["Sales Person Name"])
                c_sp_mob = get_val(["Sales Person Mobile"])
                c_sp_mail = get_val(["Sales Person Email"])
                c_pay = get_val(["Payment Terms"])

                # Logic: If Country is India -> Domestic, Else -> Export
                if country.lower() == "india":
                    # Domestic Row
                    row_data = [
                        c_name, c_bill, 
                        get_val(["State"]), get_val(["GST No", "GST"]), 
                        c_mob, 
                        get_val(["Ship Site Name", "Site"]), get_val(["Shipping Address", "Ship Address"]),
                        c_sp, c_sp_mob, c_sp_mail, c_pay
                    ]
                    ws_dom.append(row_data)
                else:
                    # Export Row
                    row_data = [
                        c_name, c_bill, country, c_mob,
                        get_val(["Ship Site Name", "Consignee Name"]), # Map Site -> Consignee
                        get_val(["Shipping Address", "Consignee Address"]), # Map Address -> Consignee Addr
                        c_sp, c_sp_mob, c_sp_mail,
                        get_val(["Port of Loading", "POL"]),
                        get_val(["Port of Discharge", "POD"]),
                        c_pay
                    ]
                    ws_exp.append(row_data)

            # Save Files
            wb_dom.save(os.path.join(DATA_DIR, "customers_domestic.xlsx"))
            wb_exp.save(os.path.join(DATA_DIR, "customers_export.xlsx"))
            logging.info("Auto-split completed successfully.")

        except Exception as e:
            logging.error(f"Failed to auto-split master file: {e}")
            messagebox.showerror("Data Error", f"Failed to process customers.xlsx\n{e}")

    @staticmethod
    def _load_customer_file(filename, is_export=False):
        """Generic loader for domestic or export customer files."""
        cust_dict = {}
        path = os.path.join(DATA_DIR, filename)
        
        if not os.path.exists(path):
            logging.warning(f"{filename} not found at: {path}")
            return cust_dict

        try:
            wb = openpyxl.load_workbook(path, data_only=True)
            ws = wb.active

            headers = {
                DataManager.normalize(cell.value): i
                for i, cell in enumerate(ws[1])
                if cell.value is not None and str(cell.value).strip() != ""
            }

            # Common Columns
            col_name = DataManager._find_col(headers, ["Customer Name", "Name"])
            col_bill = DataManager._find_col(headers, ["Billing Address"])
            col_mobile = DataManager._find_col(headers, ["Mobile", "Phone"])
            col_sp_name = DataManager._find_col(headers, ["Sales Person Name"])
            col_sp_mob = DataManager._find_col(headers, ["Sales Person Mobile"])
            col_sp_mail = DataManager._find_col(headers, ["Sales Person Email"])
            col_pay = DataManager._find_col(headers, ["Payment Terms", "Payment Term"])
            
            # Conditional Columns
            if is_export:
                col_state = None
                col_gst = None
                col_country = DataManager._find_col(headers, ["Destination Country", "Country"])
                col_site = DataManager._find_col(headers, ["Consignee Name", "Ship Site Name"])
                col_shipaddr = DataManager._find_col(headers, ["Consignee Address", "Shipping Address"])
                col_pol = DataManager._find_col(headers, ["Port of Loading", "POL"])
                col_pod = DataManager._find_col(headers, ["Port of Discharge", "POD"])
            else:
                col_state = DataManager._find_col(headers, ["State"])
                col_gst = DataManager._find_col(headers, ["GST No", "GST"])
                col_country = None # Default to India
                col_site = DataManager._find_col(headers, ["Ship Site Name", "Site"])
                col_shipaddr = DataManager._find_col(headers, ["Shipping Address", "Ship Address"])
                col_pol = None
                col_pod = None

            if col_name is None:
                logging.error(f"{filename}: 'Customer Name' column not found.")
                return cust_dict

            for row in ws.iter_rows(min_row=2, values_only=True):
                raw = row[col_name] if col_name < len(row) else None
                name = str(raw).strip() if raw is not None else ""
                if not name: continue

                def s(idx, default=""):
                    if idx is None or idx >= len(row) or row[idx] is None:
                        return default
                    return str(row[idx])

                if name not in cust_dict:
                    cust_dict[name] = []

                entry = {
                    "billing_addr": s(col_bill, ""),
                    "mobile": s(col_mobile, ""),
                    "sp_name": s(col_sp_name, ""),
                    "sp_mob": s(col_sp_mob, ""),
                    "sp_mail": s(col_sp_mail, ""),
                    "payment_terms": s(col_pay, ""),
                    "site_name": s(col_site, ""),
                    "shipping_addr": s(col_shipaddr, "")
                }

                if is_export:
                    entry["country"] = s(col_country, "")
                    entry["pol"] = s(col_pol, "")
                    entry["pod"] = s(col_pod, "")
                    entry["gst"] = ""
                    entry["state"] = ""
                else:
                    entry["gst"] = s(col_gst, "")
                    entry["state"] = s(col_state, "")
                    entry["country"] = "India"
                    entry["pol"] = ""
                    entry["pod"] = ""

                cust_dict[name].append(entry)

            logging.info(f"Loaded {len(cust_dict)} customers from {filename}")
        except Exception as e:
            logging.error(f"Error loading {filename}: {e}")
        
        return cust_dict

    @staticmethod
    def load_data():
        DataManager.products = []
        DataManager.products_map = {}
        
        # 1. Load Products
        p_path = os.path.join(DATA_DIR, "products.xlsx")
        if os.path.exists(p_path):
            try:
                wb = openpyxl.load_workbook(p_path, data_only=True)
                ws = wb.active
                headers = {
                    DataManager.normalize(cell.value): i
                    for i, cell in enumerate(ws[1])
                    if cell.value is not None and str(cell.value).strip() != ""
                }

                col_quality = DataManager._find_col(headers, ["Quality"])
                col_desc = DataManager._find_col(headers, ["Description", "Desc"])
                col_design = DataManager._find_col(headers, ["Design No", "Design"])
                col_shade = DataManager._find_col(headers, ["Shade No", "Shade"])
                col_hsn = DataManager._find_col(headers, ["HSN Code", "HSN"])
                col_unit = DataManager._find_col(headers, ["Unit"])
                col_inr = DataManager._find_col(headers, ["Base Rate", "Rate INR", "INR"])
                col_usd = DataManager._find_col(headers, ["Rate USD", "USD Rate", "USD"])
                col_gbp = DataManager._find_col(headers, ["Rate GBP", "POUND", "GBP"])

                if col_quality is not None:
                    for row in ws.iter_rows(min_row=2, values_only=True):
                        q = row[col_quality] if col_quality < len(row) else None
                        if q is None or str(q).strip() == "": continue

                        def s(idx, default=""): return str(row[idx]) if idx is not None and idx < len(row) and row[idx] is not None else default
                        def f(idx, default=0.0):
                            if idx is None or idx >= len(row) or row[idx] is None: return default
                            try: return float(row[idx])
                            except: return default

                        prod_obj = {
                            "quality": str(q).strip(),
                            "desc": s(col_desc, ""),
                            "design": s(col_design, ""),
                            "shade": s(col_shade, ""),
                            "hsn": s(col_hsn, ""),
                            "unit": s(col_unit, "MTR"),
                            "rate_inr": f(col_inr, 0.0),
                            "rate_usd": f(col_usd, 0.0),
                            "rate_gbp": f(col_gbp, 0.0),
                        }
                        DataManager.products.append(prod_obj)
                        if prod_obj["quality"] not in DataManager.products_map:
                            DataManager.products_map[prod_obj["quality"]] = []
                        DataManager.products_map[prod_obj["quality"]].append(prod_obj)

                logging.info(f"Loaded products: {len(DataManager.products)}")
            except Exception as e:
                logging.error(f"Products Load Error: {e}")
        else:
            logging.warning(f"products.xlsx not found")

        # 2. AUTO-SPLIT LOGIC (FORCED)
        master_path = os.path.join(DATA_DIR, "customers.xlsx")
        if os.path.exists(master_path):
            DataManager._auto_split_master_file(master_path)
        # 3. Load Files
        DataManager.customers_domestic = DataManager._load_customer_file("customers_domestic.xlsx", is_export=False)
        DataManager.customers_export = DataManager._load_customer_file("customers_export.xlsx", is_export=True)


# =================================================
# DOC GENERATOR
# =================================================
class DocGenerator:
    @staticmethod
    def _add_page_borders(section):
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for border_name in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')
            pgBorders.append(border)
        sectPr.append(pgBorders)

    @staticmethod
    def _add_bottom_border(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    @staticmethod
    def generate(data: dict):
        doc = Document()
        
        # Page Setup
        section = doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(3.5)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        DocGenerator._add_page_borders(section)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        # Header
        head = doc.add_paragraph("SALES INDENT")
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        head.runs[0].bold = True
        head.runs[0].font.size = Pt(16)
        
        doc.add_paragraph("")
        comp = COMPANY_DB.get(data["company"])
        if not comp: raise ValueError("Invalid Company")
        
        currency = data.get("currency", "INR")

        h = doc.add_table(rows=1, cols=2)
        h.autofit = False
        h.columns[0].width = Inches(4.5)
        h.columns[1].width = Inches(2.5)
        
        c1 = h.cell(0, 0)
        p = c1.paragraphs[0]
        p.add_run("SUPPLIER:\n").bold = True
        p.add_run(data["company"].upper() + "\n").bold = True
        p.add_run(comp["address"])

        c2 = h.cell(0, 1)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.add_run(f"GSTIN: {comp['gstin']}\n").bold = True
        
        # Add IEC Code here if Export (always from company DB, per-company)
        if data["txn_type"] == "Export" and comp.get("iec"):
            p2.add_run(f"IEC: {comp['iec']}\n").bold = True
             
        p2.add_run(f"Ph: {comp['phone']}")
        
        doc.add_paragraph("")

        # References
        ref = doc.add_table(rows=2, cols=2)
        ref.style = "Table Grid"
        ref.cell(0, 0).text = f"Our Ref: {data['our_ref']}\nDate: {data['date']}"
        ref.cell(0, 1).text = f"Buyer Ref: {data['buyer_ref']}\nOrder Ref: {data['ord_ref']}"
        
        buyer_text = f"Buyer:\n{data['buyer_name']}\n{data['bill_addr']}"
        if data["txn_type"] == "Domestic":
            buyer_text += f"\nGST: {data['buyer_gst']}\nState: {data['buyer_state']}"
        else:
            buyer_text += f"\nCountry: {data.get('country_dest', '')}"
            
        ref.cell(1, 0).text = buyer_text
        ref.cell(1, 1).text = f"Consignee:\n{data['ship_site']}\n{data['ship_addr']}\nContact: {data['ship_contact']}"

        # Sales Rep
        if data.get("sales_name", "").strip():
            doc.add_paragraph("")
            sp_table = doc.add_table(rows=1, cols=2)
            sp_table.style = "Table Grid"
            sp_table.autofit = False
            sp_table.columns[0].width = Inches(2.0)
            sp_table.columns[1].width = Inches(5.0)
            sp_table.cell(0, 0).text = "Sales Representative"
            sp_table.cell(0, 0).paragraphs[0].runs[0].bold = True
            sp_table.cell(0, 1).text = f"{data.get('sales_name','')} | {data.get('sales_mob','')} | {data.get('sales_mail','')}"

        # Export Logistics
        if data.get("txn_type") == "Export":
            doc.add_paragraph("")
            e = doc.add_table(rows=4, cols=4)
            e.style = "Table Grid"
            e.cell(0, 0).text = "Country of Origin"; e.cell(0, 1).text = data.get("country_origin", "India")
            e.cell(0, 2).text = "Country of Destination"; e.cell(0, 3).text = data.get("country_dest", "")
            e.cell(1, 0).text = "Port of Loading"; e.cell(1, 1).text = data.get("port_load", "")
            e.cell(1, 2).text = "Port of Discharge"; e.cell(1, 3).text = data.get("port_dis", "")
            e.cell(2, 0).text = "IncoTerms"; e.cell(2, 1).text = data.get("incoterm", "")
            e.cell(2, 2).text = "Shipping Date"; e.cell(2, 3).text = data.get("shipping_date", "")
            e.rows[3].cells[0].merge(e.rows[3].cells[3])
            
            days = data.get('validity_days', 30)
            e.rows[3].cells[0].text = f"Indent is valid for {days} days."
            doc.add_paragraph("")

        # Items
        cols = ["S.No", "Quality", "Description", "Design", "Shade", "Item Code", "HSN", "Qty", "Unit", f"Rate ({currency})", f"Amount ({currency})"]
        t = doc.add_table(rows=1, cols=len(cols))
        t.style = "Table Grid"
        for i, col in enumerate(cols):
            t.rows[0].cells[i].text = col
            t.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        for idx, item in enumerate(data["items"], 1):
            r = t.add_row().cells
            r[0].text = str(idx)
            r[1].text = str(item.get("quality", ""))
            r[2].text = str(item.get("desc", ""))
            r[3].text = str(item.get("design", ""))
            r[4].text = str(item.get("shade", ""))
            ref_txt = str(item.get("buyer_ref", "")).strip()
            r[5].text = "" if ref_txt == ":" else ref_txt
            r[6].text = str(item.get("hsn", ""))
            r[7].text = f"{item.get('qty', 0):,.2f}"
            r[8].text = str(item.get("unit", ""))
            r[9].text = f"{float(item.get('rate', 0.0)):,.2f}"
            r[10].text = f"{float(item.get('amount', 0.0)):,.2f}"

        subtotal = float(data.get("subtotal", 0.0))
        grand_total = subtotal

        def add_fin(label, value):
            rr = t.add_row().cells
            rr[0].merge(rr[9])
            rr[0].text = label
            rr[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rr[10].text = f"{value:,.2f}"

        add_fin(f"Sub-Total ({currency})", subtotal)

        # TAX LOGIC (Domestic Only)
        if data.get("txn_type") == "Domestic" and currency == "INR":
            buyer_state = str(data.get("buyer_state", "")).lower()
            buyer_gst = str(data.get("buyer_gst", ""))
            is_guj = ("gujarat" in buyer_state) or buyer_gst.startswith("24")

            if is_guj:
                tax = subtotal * GST_RATE_CGST
                add_fin(f"CGST ({GST_RATE_CGST*100}%)", tax)
                add_fin(f"SGST ({GST_RATE_SGST*100}%)", tax)
                grand_total += tax * 2
            else:
                tax = subtotal * GST_RATE_IGST
                add_fin(f"IGST ({GST_RATE_IGST*100}%)", tax)
                grand_total += tax

        # Totals
        gt = t.add_row().cells
        gt[0].merge(gt[9])
        
        p = gt[0].paragraphs[0]
        p.text = "" 
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"Amount in Words: {num2words(grand_total, lang='en').title()} {currency} Only\n\nGRAND TOTAL ({currency})").bold = True
        
        gt[10].text = ""
        p2 = gt[10].paragraphs[0]
        p2.add_run(f"\n\n{grand_total:,.2f}").bold = True
        
        doc.add_paragraph("")

        # Terms
        for title, key in [("Sampling Requirements:", "sampling"), ("Packaging Requirements:", "packaging")]:
            val = (data.get(key) or "").strip()
            if val:
                p = doc.add_paragraph()
                p.add_run(title).bold = True
                for line in val.split("\n"):
                    if line.strip(): doc.add_paragraph(line.strip(), style="List Paragraph")

        p = doc.add_paragraph()
        p.add_run("Notes:").bold = True
        if data.get("terms"):
            for line in data["terms"].split("\n"):
                if line.strip(): doc.add_paragraph(line.strip(), style="List Paragraph")

        doc.add_paragraph("Any Dispute regarding the goods must be raised with the respective Marketing Person within 30 Days from the date of receipt of goods.", style="List Paragraph")
        doc.add_paragraph().add_run(f"Payment Terms: {data.get('payment_terms', '')}").bold = True

        # Bank
        doc.add_paragraph("\nBANK DETAILS:").runs[0].bold = True
        b = comp["bank_details"]
        bank_cols = ["Account Holder", "Bank", "Branch", "Acct", "IFSC"]
        if data.get("txn_type") == "Export": bank_cols.append("SWIFT")
        
        bt = doc.add_table(rows=2, cols=len(bank_cols))
        bt.style = "Table Grid"
        for i, k in enumerate(bank_cols):
            bt.cell(0, i).text = k
            bt.cell(0, i).paragraphs[0].runs[0].bold = True
            bt.cell(1, i).text = str(b.get(k, ""))

        if data.get("txn_type") == "Export":
            doc.add_paragraph("\nDocuments Required:").bold = True
            for d in ["Commercial Invoice - 3 Copies", "Bill of Lading - 3 Copies", "Packing List - 3 Copies", "Certificate of Origin - 3 Copies"]:
                doc.add_paragraph(d, style="List Bullet")

        # Signature
        doc.add_paragraph("")
        DocGenerator._add_bottom_border(doc.add_paragraph())
        sig = doc.add_table(rows=2, cols=3)
        sig.autofit = True
        sig.cell(0, 0).text = f"For, {data['company']}"
        sig.cell(0, 1).text = data["company"]
        sig.cell(0, 2).text = f"For, {data['buyer_name']}"
        sig.cell(1, 0).text = "\n\n\n\nAuthorised Signatory"
        sig.cell(1, 1).text = "\n\n\n\nDirector"
        sig.cell(1, 2).text = "\n\n\n\nSignature & Seal of Company"

        # Save
        ts = datetime.now().strftime("%d%b_%H%M")
        filename = f"{data['company'][:4]}_{data['our_ref']}_{ts}_Indent.docx"
        clean_name = "".join([c for c in filename if c.isalpha() or c.isdigit() or c in (" ", "-", "_", ".")]).rstrip()
        out_path = os.path.join(OUTPUT_DIR, clean_name)
        
        try:
            doc.save(out_path)
            return out_path
        except PermissionError:
            raise PermissionError(f"File is open in Word. Please close '{clean_name}' and try again.")


# =================================================
# NEW PRODUCT SELECTOR POPUP
# =================================================
class ProductPopup(ctk.CTkToplevel):
    def __init__(self, master, quality, items, currency, callback):
        super().__init__(master)
        self.title(f"Select Items for Quality: {quality}")
        self.geometry("900x600")
        self.callback = callback
        self.currency = currency
        self.quality = quality
        
        self.transient(master)
        self.grab_set()
        
        top_frame = ctk.CTkFrame(self, fg_color="transparent")
        top_frame.pack(fill="x", padx=12, pady=12)
        ctk.CTkLabel(top_frame, text=f"Product Selection: {quality}", font=("Arial", 15, "bold")).pack(side="left")
        ctk.CTkButton(top_frame, text="ADD SELECTED TO CART", command=self._submit, fg_color="#2E7D32", hover_color="#1B5E20").pack(side="right")

        header_frame = ctk.CTkFrame(self, fg_color="#E8E8E8", height=36, corner_radius=0)
        header_frame.pack(fill="x", padx=12)
        
        ctk.CTkLabel(header_frame, text="Select", width=50, text_color="#37474F", font=("Arial", 11, "bold")).grid(row=0, column=0, padx=8, pady=6)
        ctk.CTkLabel(header_frame, text="Design/Shade", width=200, anchor="w", text_color="#37474F", font=("Arial", 11, "bold")).grid(row=0, column=1, padx=8, pady=6)
        ctk.CTkLabel(header_frame, text="Description", width=200, anchor="w", text_color="#37474F", font=("Arial", 11, "bold")).grid(row=0, column=2, padx=8, pady=6)
        ctk.CTkLabel(header_frame, text=f"Rate ({currency})", width=100, text_color="#37474F", font=("Arial", 11, "bold")).grid(row=0, column=3, padx=8, pady=6)
        ctk.CTkLabel(header_frame, text="Quantity", width=100, text_color="#37474F", font=("Arial", 11, "bold")).grid(row=0, column=4, padx=8, pady=6)

        self.scroll = ctk.CTkScrollableFrame(self, fg_color="transparent")
        self.scroll.pack(fill="both", expand=True, padx=12, pady=8)
        
        self.row_entries = []

        for i, item in enumerate(items):
            row_frame = ctk.CTkFrame(self.scroll)
            row_frame.pack(fill="x", pady=2)
            
            var = ctk.BooleanVar()
            cb = ctk.CTkCheckBox(row_frame, text="", variable=var, width=50)
            cb.grid(row=0, column=0, padx=5)
            
            ref_str = f"{item['design']} / {item['shade']}"
            ctk.CTkLabel(row_frame, text=ref_str, width=200, anchor="w").grid(row=0, column=1, padx=5)
            ctk.CTkLabel(row_frame, text=item['desc'][:30], width=200, anchor="w").grid(row=0, column=2, padx=5)
            
            rate_val = item['rate_usd'] if currency == "USD" else (item['rate_gbp'] if currency == "GBP" else item['rate_inr'])
            e_rate = ctk.CTkEntry(row_frame, width=100)
            e_rate.insert(0, str(rate_val))
            e_rate.grid(row=0, column=3, padx=5)
            
            e_qty = ctk.CTkEntry(row_frame, width=100, placeholder_text="0.0")
            e_qty.grid(row=0, column=4, padx=5)
            
            self.row_entries.append({
                "var": var,
                "item": item,
                "e_rate": e_rate,
                "e_qty": e_qty,
                "ref_str": ref_str
            })

    def _submit(self):
        selected_items = []
        for entry in self.row_entries:
            if entry["var"].get():
                try:
                    qty = float(entry["e_qty"].get())
                    rate = float(entry["e_rate"].get())
                    if qty <= 0: continue
                    
                    original_item = entry["item"]
                    cart_item = {
                        "quality": original_item["quality"],
                        "desc": original_item["desc"],
                        "design": original_item["design"],
                        "shade": original_item["shade"],
                        "hsn": original_item["hsn"],
                        "unit": original_item["unit"],
                        "qty": qty,
                        "rate": rate,
                        "amount": qty * rate,
                        "buyer_ref": entry["ref_str"]
                    }
                    selected_items.append(cart_item)
                except ValueError:
                    continue

        if not selected_items:
            messagebox.showwarning("No Items", "No items selected or valid quantities entered.")
            return

        self.callback(selected_items)
        self.destroy()


# =================================================
# UI APP
# =================================================
class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Gujarat Flotex Sales Indent v16.1")
        self.geometry("1500x950")
        self.cart = []
        self._data_loaded = False
        self._init_sidebar()
        self._init_main_area()
        self._bind_keys()
        self._toggle_export_ui("Domestic")
        self._reset_form_only(show_message=False)
        # Defer data load to background so window appears immediately and stays responsive
        self._loading_label = ctk.CTkLabel(self.sidebar, text="Loading data...", text_color=LABEL_GRAY, font=("Arial", 11))
        self._loading_label.pack(pady=(8, 0))
        self.after(50, self._start_data_load)
        self.update_idletasks()

    def _bind_keys(self):
        def _on_enter(event):
            if event.widget.winfo_class() == "Text": return
            if hasattr(event.widget, "tk_focusNext"):
                target = event.widget.tk_focusNext()
                if target: target.focus(); return "break"
        self.bind("<Return>", _on_enter)

    def _start_data_load(self):
        """Run data loading in a background thread so UI stays responsive."""
        def load():
            try:
                DataManager.load_data()
            except Exception as e:
                logging.error(f"Data load error: {e}")
            # Defer UI update by one event-loop tick so the window can paint and stay responsive (avoids "Not responding")
            self.after(0, lambda: self.after(1, self._on_data_loaded))
        threading.Thread(target=load, daemon=True).start()

    def _on_data_loaded(self):
        """Called on main thread after background load finishes; refresh dropdowns."""
        self._data_loaded = True
        if hasattr(self, "_loading_label") and self._loading_label.winfo_exists():
            self._loading_label.configure(text="Ready")
        self._on_txn_type_change(self.txn_var.get())
        self.prod_search.configure(values=[PH_PRODUCT] + list(DataManager.products_map.keys()))

    def _init_sidebar(self):
        self.sidebar = ctk.CTkFrame(self, width=280, corner_radius=0, fg_color="#F5F5F5")
        self.sidebar.pack(side="left", fill="y")

        ctk.CTkButton(self.sidebar, text="REFRESH DATA", fg_color="#D32F2F", hover_color="#B71C1C", command=self._reset_and_refresh).pack(fill="x", padx=16, pady=(16, 8))
        ctk.CTkButton(self.sidebar, text="NEW INDENT", fg_color="#455A64", hover_color="#37474F", command=self._new_pi).pack(fill="x", padx=16, pady=(0, 16))
        
        ctk.CTkLabel(self.sidebar, text="SETTINGS", font=("Arial", 14, "bold"), text_color=LABEL_GRAY).pack(pady=(8, 4))
        self.comp_var = ctk.StringVar(value=PH_COMPANY)
        self.comp_menu = ctk.CTkOptionMenu(self.sidebar, variable=self.comp_var, values=[PH_COMPANY] + list(COMPANY_DB.keys()), command=self._on_company_change)
        self.comp_menu.pack(fill="x", padx=16, pady=4)

        self.txn_var = ctk.StringVar(value="Domestic")
        self.txn_seg = ctk.CTkSegmentedButton(self.sidebar, variable=self.txn_var, values=["Domestic", "Export"], command=self._on_txn_type_change)
        self.txn_seg.pack(fill="x", padx=16, pady=8)

        self.currency_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        ctk.CTkLabel(self.currency_frame, text="Export Currency", anchor="w", text_color=LABEL_GRAY).pack(fill="x", pady=(8, 0))
        self.currency_var = ctk.StringVar(value=EXPORT_CURRENCIES[0])
        self.currency_menu = ctk.CTkOptionMenu(self.currency_frame, variable=self.currency_var, values=EXPORT_CURRENCIES, command=lambda _: self._on_currency_change())
        self.currency_menu.pack(fill="x", pady=4)

        ctk.CTkLabel(self.sidebar, text="REFERENCES", font=("Arial", 12, "bold"), anchor="w", text_color=LABEL_GRAY).pack(fill="x", padx=16, pady=(12, 4))
        for lbl, entry_attr in [("Our Ref No *", "ref_entry"), ("Buyer PO No *", "buyer_ref"), ("Order Ref", "ord_ref")]:
            ctk.CTkLabel(self.sidebar, text=lbl, text_color=LABEL_GRAY, anchor="w").pack(fill="x", padx=16)
            setattr(self, entry_attr, ctk.CTkEntry(self.sidebar))
            getattr(self, entry_attr).pack(fill="x", padx=16, pady=(0, 6))

        ctk.CTkLabel(self.sidebar, text="Indent Date", text_color=LABEL_GRAY, anchor="w").pack(fill="x", padx=16)
        self.date_pick = DateEntry(self.sidebar, width=12)
        self.date_pick.pack(padx=16, pady=(0, 8), anchor="w")

    def _init_main_area(self):
        self.main_scroll = ctk.CTkScrollableFrame(self, fg_color="transparent")
        self.main_scroll.pack(side="right", fill="both", expand=True, padx=16, pady=12)

        # ------------------ BUYER CARD ------------------
        self.buyer_frame = ctk.CTkFrame(self.main_scroll, fg_color="white", corner_radius=CARD_RADIUS, border_width=1, border_color=CARD_BORDER)
        self.buyer_frame.pack(fill="x", pady=(0, 10), padx=4)
        ctk.CTkLabel(self.buyer_frame, text="BUYER & SALES DETAILS", font=SECTION_HEADER_FONT).grid(row=0, column=0, columnspan=3, sticky="w", padx=CARD_PAD, pady=(CARD_PAD, 8))
        
        # Row 1: Labels
        ctk.CTkLabel(self.buyer_frame, text="Buyer *", text_color=LABEL_GRAY).grid(row=1, column=0, sticky="w", padx=CARD_PAD)
        ctk.CTkLabel(self.buyer_frame, text="Consignee", text_color=LABEL_GRAY).grid(row=1, column=1, sticky="w", padx=CARD_PAD)
        
        # Row 2: Inputs
        self.buyer_combo = ctk.CTkComboBox(self.buyer_frame, values=[PH_BUYER], command=self._on_buyer, width=300)
        self.buyer_combo.set(PH_BUYER)
        self.buyer_combo.grid(row=2, column=0, padx=CARD_PAD, pady=4)
        
        self.ship_var = ctk.StringVar(value=PH_SHIP)
        self.ship_combo = ctk.CTkOptionMenu(self.buyer_frame, variable=self.ship_var, values=[PH_SHIP], width=300)
        self.ship_combo.grid(row=2, column=1, padx=CARD_PAD, pady=4)

        # Row 3: Sales Labels
        ctk.CTkLabel(self.buyer_frame, text="Sales Name", text_color=LABEL_GRAY).grid(row=3, column=0, sticky="w", padx=CARD_PAD, pady=(10, 0))
        ctk.CTkLabel(self.buyer_frame, text="Sales Mobile", text_color=LABEL_GRAY).grid(row=3, column=1, sticky="w", padx=CARD_PAD, pady=(10, 0))
        ctk.CTkLabel(self.buyer_frame, text="Sales Email", text_color=LABEL_GRAY).grid(row=3, column=2, sticky="w", padx=CARD_PAD, pady=(10, 0))

        # Row 4: Sales Inputs
        self.sales_name = ctk.CTkEntry(self.buyer_frame, placeholder_text="Sales Name", width=300)
        self.sales_name.grid(row=4, column=0, padx=CARD_PAD, pady=(0, CARD_PAD))
        self.sales_mob = ctk.CTkEntry(self.buyer_frame, placeholder_text="Mobile", width=140)
        self.sales_mob.grid(row=4, column=1, padx=CARD_PAD, sticky="w", pady=(0, CARD_PAD))
        self.sales_mail = ctk.CTkEntry(self.buyer_frame, placeholder_text="Email", width=200)
        self.sales_mail.grid(row=4, column=2, padx=CARD_PAD, sticky="w", pady=(0, CARD_PAD))

        # ------------------ EXPORT/LOGISTICS CARD ------------------
        self.export_frame = ctk.CTkFrame(self.main_scroll, fg_color="white", corner_radius=CARD_RADIUS, border_width=1, border_color=CARD_BORDER)
        ctk.CTkLabel(self.export_frame, text="LOGISTICS & SHIPMENT", font=SECTION_HEADER_FONT).grid(row=0, column=0, columnspan=3, sticky="w", padx=CARD_PAD, pady=(CARD_PAD, 8))
        
        ctk.CTkLabel(self.export_frame, text="Incoterm", text_color=LABEL_GRAY).grid(row=1, column=0, sticky="w", padx=CARD_PAD)
        ctk.CTkLabel(self.export_frame, text="Country of Origin", text_color=LABEL_GRAY).grid(row=1, column=1, sticky="w", padx=CARD_PAD)
        ctk.CTkLabel(self.export_frame, text="Country of Destination", text_color=LABEL_GRAY).grid(row=1, column=2, sticky="w", padx=CARD_PAD)

        self.incoterm = ctk.CTkComboBox(self.export_frame, values=[""] + INCO_TERMS)
        self.incoterm.set("")
        self.incoterm.grid(row=2, column=0, padx=CARD_PAD, pady=4)
        self.origin_country = ctk.CTkEntry(self.export_frame)
        self.origin_country.grid(row=2, column=1, padx=CARD_PAD, pady=4)
        self.dest_country = ctk.CTkEntry(self.export_frame) 
        self.dest_country.grid(row=2, column=2, padx=CARD_PAD, pady=4)

        ctk.CTkLabel(self.export_frame, text="Port of Loading", text_color=LABEL_GRAY).grid(row=3, column=0, sticky="w", padx=CARD_PAD, pady=(10, 0))
        ctk.CTkLabel(self.export_frame, text="Port of Discharge", text_color=LABEL_GRAY).grid(row=3, column=1, sticky="w", padx=CARD_PAD, pady=(10, 0))
        ctk.CTkLabel(self.export_frame, text="Shipping Date", text_color=LABEL_GRAY).grid(row=3, column=2, sticky="w", padx=CARD_PAD, pady=(10, 0))

        self.port_load = ctk.CTkEntry(self.export_frame)
        self.port_load.grid(row=4, column=0, padx=CARD_PAD, pady=4)
        self.port_dis = ctk.CTkEntry(self.export_frame)
        self.port_dis.grid(row=4, column=1, padx=CARD_PAD, pady=4)
        self.shipping_date = DateEntry(self.export_frame, width=12)
        self.shipping_date.grid(row=4, column=2, padx=CARD_PAD, pady=4, sticky="w")
        
        ctk.CTkLabel(self.export_frame, text="Validity (Days)", text_color=LABEL_GRAY).grid(row=5, column=0, sticky="w", padx=CARD_PAD, pady=(10, 0))
        self.validity_days = ctk.CTkEntry(self.export_frame, width=80)
        self.validity_days.insert(0, "30")
        self.validity_days.grid(row=6, column=0, padx=CARD_PAD, pady=(0, CARD_PAD), sticky="w")
        # IEC is taken from company DB automatically (read-only display when company selected)
        self.iec_label = ctk.CTkLabel(self.export_frame, text="IEC: (select company)", text_color=LABEL_GRAY, anchor="w")
        self.iec_label.grid(row=5, column=1, rowspan=2, sticky="w", padx=CARD_PAD, pady=(10, CARD_PAD))

        # ------------------ PRODUCT CARD ------------------
        self.prod_frame = ctk.CTkFrame(self.main_scroll, fg_color="white", corner_radius=CARD_RADIUS, border_width=1, border_color=CARD_BORDER)
        self.prod_frame.pack(fill="x", pady=(0, 10), padx=4)
        ctk.CTkLabel(self.prod_frame, text="PRODUCT ITEMS", font=SECTION_HEADER_FONT).grid(row=0, column=0, columnspan=3, sticky="w", padx=CARD_PAD, pady=(CARD_PAD, 8))
        
        ctk.CTkLabel(self.prod_frame, text="Select Quality to Open Items:", text_color=LABEL_GRAY).grid(row=1, column=0, sticky="w", padx=CARD_PAD)
        
        self.prod_search = ctk.CTkComboBox(self.prod_frame, values=[PH_PRODUCT] + list(DataManager.products_map.keys()), command=self._open_product_popup, width=350)
        self.prod_search.set(PH_PRODUCT)
        self.prod_search.grid(row=2, column=0, padx=CARD_PAD, pady=(0, CARD_PAD))
        
        ctk.CTkButton(self.prod_frame, text="DELETE LAST ITEM", command=self._delete_last_item, width=150, fg_color="#D32F2F", hover_color="#B71C1C").grid(row=2, column=1, padx=CARD_PAD)
        ctk.CTkButton(self.prod_frame, text="CLEAR CART", command=self._clear_cart, width=120, fg_color="#757575", hover_color="#616161").grid(row=2, column=2, padx=CARD_PAD)

        # Cart Table
        self.table_frame = ctk.CTkFrame(self.main_scroll, fg_color="#FAFAFA", corner_radius=8, border_width=1, border_color=CARD_BORDER)
        self.table_frame.pack(fill="x", padx=4, pady=(0, 8))
        self._render_table_header()

        # Terms
        self.term_frame = ctk.CTkFrame(self.main_scroll, fg_color="white", corner_radius=CARD_RADIUS, border_width=1, border_color=CARD_BORDER)
        self.term_frame.pack(fill="x", pady=(0, 10), padx=4)
        self.txt_smp = self._create_box("1. Sampling Requirements:")
        self.txt_pkg = self._create_box("2. Packaging Requirements:")
        self.txt_trm = self._create_box("3. Notes:")
        
        pf = ctk.CTkFrame(self.term_frame, fg_color="transparent")
        pf.pack(fill="x", padx=CARD_PAD, pady=(4, CARD_PAD))
        ctk.CTkLabel(pf, text="Payment Terms:", text_color=LABEL_GRAY).pack(side="left")
        
        self.pay_combo = ctk.CTkComboBox(pf, values=[""] + PAYMENT_TERMS, width=400)
        self.pay_combo.pack(side="left", padx=8)

        # Action Buttons
        btn_frame = ctk.CTkFrame(self.main_scroll, fg_color="transparent")
        btn_frame.pack(fill="x", pady=16, padx=4)
        ctk.CTkButton(btn_frame, text="GENERATE WORD", command=self._generate_word, height=48, font=("Arial", 14, "bold"), fg_color="#0288D1", hover_color="#0277BD").pack(side="left", fill="x", expand=True, padx=(0, 8))
        ctk.CTkButton(btn_frame, text="GENERATE PDF", command=self._generate_pdf, height=48, font=("Arial", 14, "bold"), fg_color="#C62828", hover_color="#B71C1C").pack(side="left", fill="x", expand=True, padx=(8, 0))

    def _create_box(self, title):
        ctk.CTkLabel(self.term_frame, text=title, font=("Arial", 11, "bold"), text_color=LABEL_GRAY).pack(anchor="w", padx=CARD_PAD, pady=(8, 0))
        box = ctk.CTkTextbox(self.term_frame, height=60, fg_color="#F5F5F5", text_color="black", border_width=1, border_color=CARD_BORDER)
        box.pack(fill="x", padx=CARD_PAD, pady=(0, 6))
        return box

    def _render_table_header(self):
        for w in self.table_frame.winfo_children(): w.destroy()
        for i, h in enumerate(["S.No", "Quality", "Ref", "Qty", "Rate", "Amount"]):
            lbl = ctk.CTkLabel(self.table_frame, text=h, font=("Arial", 11, "bold"), width=120, text_color="#37474F", fg_color="#E8E8E8")
            lbl.grid(row=0, column=i, padx=8, pady=6, sticky="ew")

    def _render_cart_rows(self):
        self._render_table_header()
        for idx, item in enumerate(self.cart, 1):
            ctk.CTkLabel(self.table_frame, text=str(idx), text_color="black").grid(row=idx, column=0, padx=8, pady=4)
            ctk.CTkLabel(self.table_frame, text=item["quality"], text_color="black").grid(row=idx, column=1, padx=8, pady=4)
            ctk.CTkLabel(self.table_frame, text=item["buyer_ref"], text_color="black").grid(row=idx, column=2, padx=8, pady=4)
            ctk.CTkLabel(self.table_frame, text=f"{item['qty']} {item['unit']}", text_color="black").grid(row=idx, column=3, padx=8, pady=4)
            ctk.CTkLabel(self.table_frame, text=f"{item['rate']:.2f}", text_color="black").grid(row=idx, column=4, padx=8, pady=4)
            ctk.CTkLabel(self.table_frame, text=f"{item['amount']:.2f}", text_color="black").grid(row=idx, column=5, padx=8, pady=4)

    def _get_invoice_currency(self):
        return self.currency_var.get() if self.txn_var.get() == "Export" else "INR"

    def _on_company_change(self, company):
        self._reset_form_only(show_message=False)
        self.comp_var.set(company)
        # Show company IEC (read-only; document uses this from COMPANY_DB)
        comp_data = COMPANY_DB.get(company, {})
        iec = comp_data.get("iec", "") or "(not set)"
        self.iec_label.configure(text=f"IEC: {iec}")

    def _on_txn_type_change(self, txn_type):
        self._toggle_export_ui(txn_type)
        self.txn_var.set(txn_type)
        
        # CRITICAL FIX: CLEAR EVERYTHING ON SWITCH
        self._clear_cart()
        self.buyer_combo.set(PH_BUYER)
        self.ship_combo.configure(values=[PH_SHIP])
        self.ship_var.set(PH_SHIP)
        
        if txn_type == "Domestic":
            buyers = list(DataManager.customers_domestic.keys())
        else:
            buyers = list(DataManager.customers_export.keys())
            
        self.buyer_combo.configure(values=[PH_BUYER] + buyers)
        
        # Re-trigger company change to refill IEC if needed
        if self.comp_var.get() != PH_COMPANY:
            self._on_company_change(self.comp_var.get())

    def _toggle_export_ui(self, val):
        if val == "Export":
            self.export_frame.pack(fill="x", pady=(0, 10), padx=4, after=self.buyer_frame)
            self.currency_frame.pack(fill="x", padx=16, after=self.txn_seg)
        else:
            self.export_frame.pack_forget()
            self.currency_frame.pack_forget()
        self._on_currency_change()

    def _on_currency_change(self):
        pass

    def _on_buyer(self, name):
        txn = self.txn_var.get()
        source_db = DataManager.customers_domestic if txn == "Domestic" else DataManager.customers_export
        
        if name not in source_db: return
        cust = source_db[name][0]
        
        # Populate Consignee Combo
        sites = [r["site_name"] for r in source_db[name] if r.get("site_name")]
        self.ship_combo.configure(values=[PH_SHIP] + (sites if sites else ["Same as Billing"]))
        
        # Sales Details
        self.sales_name.delete(0, "end"); self.sales_name.insert(0, cust.get("sp_name", ""))
        self.sales_mob.delete(0, "end"); self.sales_mob.insert(0, cust.get("sp_mob", ""))
        self.sales_mail.delete(0, "end"); self.sales_mail.insert(0, cust.get("sp_mail", ""))

        # Auto-Fill Logistics from Master
        if txn == "Export":
            self.port_load.delete(0, "end"); self.port_load.insert(0, cust.get("pol", ""))
            self.port_dis.delete(0, "end"); self.port_dis.insert(0, cust.get("pod", ""))
            self.dest_country.delete(0, "end"); self.dest_country.insert(0, cust.get("country", ""))
        
        self.pay_combo.set(cust.get("payment_terms", ""))

    def _open_product_popup(self, quality):
        if quality == PH_PRODUCT: return
        
        items = DataManager.products_map.get(quality, [])
        if not items:
            messagebox.showerror("Error", "No items found for this Quality.")
            return

        currency = self._get_invoice_currency()
        ProductPopup(self, quality, items, currency, self._add_items_from_popup)
        self.prod_search.set(PH_PRODUCT)

    def _add_items_from_popup(self, item_list):
        self.cart.extend(item_list)
        self._render_cart_rows()

    def _delete_last_item(self):
        if self.cart: self.cart.pop(); self._render_cart_rows()

    def _clear_cart(self):
        self.cart = []; self._render_cart_rows()

    def _reset_form_only(self, show_message=True):
        self.buyer_combo.set(PH_BUYER)
        self.prod_search.set(PH_PRODUCT)
        self.ship_combo.configure(values=[PH_SHIP]); self.ship_var.set(PH_SHIP)
        self.cart = []
        self._render_cart_rows()
        for e in [self.ref_entry, self.buyer_ref, self.ord_ref, self.sales_name, self.sales_mob, self.sales_mail, self.origin_country, self.dest_country, self.port_load, self.port_dis]:
            e.delete(0, "end")
        self.origin_country.insert(0, "India")
        self.validity_days.delete(0, "end"); self.validity_days.insert(0, "30")
        for t in [self.txt_smp, self.txt_pkg, self.txt_trm]: t.delete("0.0", "end")
        self.pay_combo.set("")
        # Refresh buyer list for current txn type only (avoid calling _on_txn_type_change -> _on_company_change recursion)
        txn = self.txn_var.get()
        buyers = list(DataManager.customers_domestic.keys()) if txn == "Domestic" else list(DataManager.customers_export.keys())
        self.buyer_combo.configure(values=[PH_BUYER] + buyers)
        comp = self.comp_var.get()
        if comp != PH_COMPANY:
            self.iec_label.configure(text=f"IEC: {COMPANY_DB.get(comp, {}).get('iec', '') or '(not set)'}")
        else:
            self.iec_label.configure(text="IEC: (select company)")
        if show_message: messagebox.showinfo("Reset", "Form Cleared")

    def _reset_and_refresh(self):
        if not messagebox.askyesno("Confirm", "Reload Data?"):
            return
        if hasattr(self, "_loading_label") and self._loading_label.winfo_exists():
            self._loading_label.configure(text="Loading...")
        def load():
            try:
                DataManager.load_data()
            except Exception as e:
                logging.error(f"Data load error: {e}")
            self.after(0, self._on_refresh_done)
        threading.Thread(target=load, daemon=True).start()

    def _on_refresh_done(self):
        self._reset_form_only(show_message=False)
        self._on_data_loaded()
        if hasattr(self, "_loading_label") and self._loading_label.winfo_exists():
            self._loading_label.configure(text="Ready")
        messagebox.showinfo("Done", "Data reloaded.")

    def _new_pi(self): self._reset_form_only()

    def _get_form_data(self):
        if self.comp_var.get() == PH_COMPANY: messagebox.showerror("Error", "Select Company"); return None
        if not self.ref_entry.get().strip() or not self.buyer_ref.get().strip(): messagebox.showerror("Error", "Refs mandatory"); return None
        if not self.cart: messagebox.showerror("Error", "Cart Empty"); return None
        
        bn = self.buyer_combo.get()
        txn = self.txn_var.get()
        source_db = DataManager.customers_domestic if txn == "Domestic" else DataManager.customers_export

        if bn == PH_BUYER or bn not in source_db:
            messagebox.showerror("Error", "Select a valid Buyer"); return None

        cust = source_db[bn][0]
        if self.ship_var.get() != PH_SHIP and self.ship_var.get() != "Same as Billing":
            cust = next((x for x in source_db[bn] if x["site_name"] == self.ship_var.get()), cust)
            
        try: validity = int(self.validity_days.get())
        except ValueError: validity = 30

        return {
            "company": self.comp_var.get(), "txn_type": self.txn_var.get(), "currency": self._get_invoice_currency(),
            "our_ref": self.ref_entry.get(), "buyer_ref": self.buyer_ref.get(), "ord_ref": self.ord_ref.get(),
            "date": self.date_pick.get(), "buyer_name": bn, 
            "bill_addr": cust.get("billing_addr", ""), 
            "buyer_gst": cust.get("gst", ""), 
            "buyer_state": cust.get("state", ""),
            "ship_site": cust.get("site_name", ""), "ship_addr": cust.get("shipping_addr", ""), "ship_contact": cust.get("mobile", ""),
            "sales_name": self.sales_name.get(), "sales_mob": self.sales_mob.get(), "sales_mail": self.sales_mail.get(),
            "incoterm": self.incoterm.get(), "country_origin": self.origin_country.get(), 
            "country_dest": self.dest_country.get(),
            "port_load": self.port_load.get(), "port_dis": self.port_dis.get(), "shipping_date": self.shipping_date.get(),
            "validity_days": validity, "items": self.cart,
            "subtotal": sum(x["amount"] for x in self.cart), "payment_terms": self.pay_combo.get(),
            "sampling": self.txt_smp.get("0.0", "end"), "packaging": self.txt_pkg.get("0.0", "end"), "terms": self.txt_trm.get("0.0", "end")
        }

    def _generate_word(self):
        data = self._get_form_data()
        if data: 
            try: open_file(DocGenerator.generate(data))
            except Exception as e: messagebox.showerror("Error", str(e))

    def _generate_pdf(self):
        if not PDF_SUPPORT: return messagebox.showerror("Error", "Install docx2pdf")
        data = self._get_form_data()
        if data:
            try:
                docx = DocGenerator.generate(data)
                pdf = docx.replace(".docx", ".pdf")
                messagebox.showinfo("Status", "Generating PDF...")
                convert_to_pdf(docx, pdf)
                open_file(pdf)
            except Exception as e: messagebox.showerror("Error", str(e))

if __name__ == "__main__":
    App().mainloop()